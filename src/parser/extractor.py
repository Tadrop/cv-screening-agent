"""
CV text extraction from PDF and DOCX attachments.

Primary PDF parser: pypdf
Fallback PDF parser: pdfplumber (handles more complex layouts, scanned PDFs signal)
DOCX parser: python-docx
"""

from __future__ import annotations

import logging
import re
from pathlib import Path
from typing import NamedTuple

import pypdf
import pdfplumber

logger = logging.getLogger(__name__)


class ExtractionResult(NamedTuple):
    text: str
    parser_used: str          # "pypdf" | "pdfplumber" | "docx"
    page_count: int
    is_image_based: bool      # True if no text layer detected (scanned PDF)


def extract_text(file_path: Path) -> ExtractionResult:
    """
    Dispatch to the correct extractor based on file extension.
    Raises ValueError for unsupported types.
    """
    suffix = file_path.suffix.lower()
    if suffix == ".pdf":
        return _extract_pdf(file_path)
    if suffix in {".docx", ".doc"}:
        return _extract_docx(file_path)
    raise ValueError(f"Unsupported file type: {suffix!r}")


# ──────────────────────────────────────────────────────────────────────────────
# PDF
# ──────────────────────────────────────────────────────────────────────────────


def _extract_pdf(path: Path) -> ExtractionResult:
    text, page_count = _try_pypdf(path)
    parser = "pypdf"

    if not text.strip():
        logger.warning("pypdf returned empty text for %s — trying pdfplumber", path.name)
        text, page_count = _try_pdfplumber(path)
        parser = "pdfplumber"

    is_image_based = not bool(text.strip())
    if is_image_based:
        logger.warning(
            "No text layer found in %s (size=%d bytes) — likely a scanned PDF",
            path.name, path.stat().st_size,
        )
        text = "[SCANNED PDF — no text layer extracted. Manual review required.]"

    return ExtractionResult(
        text=_normalise(text),
        parser_used=parser,
        page_count=page_count,
        is_image_based=is_image_based,
    )


def _try_pypdf(path: Path) -> tuple[str, int]:
    try:
        reader = pypdf.PdfReader(str(path))
        pages = [page.extract_text() or "" for page in reader.pages]
        return "\n".join(pages), len(pages)
    except Exception as exc:
        logger.warning("pypdf failed on %s: %s", path.name, exc)
        return "", 0


def _try_pdfplumber(path: Path) -> tuple[str, int]:
    try:
        with pdfplumber.open(str(path)) as pdf:
            pages = [p.extract_text() or "" for p in pdf.pages]
        return "\n".join(pages), len(pages)
    except Exception as exc:
        logger.warning("pdfplumber failed on %s: %s", path.name, exc)
        return "", 0


# ──────────────────────────────────────────────────────────────────────────────
# DOCX
# ──────────────────────────────────────────────────────────────────────────────


def _extract_docx(path: Path) -> ExtractionResult:
    try:
        from docx import Document  # python-docx; imported lazily to keep the module light

        doc = Document(str(path))
        paragraphs = [p.text for p in doc.paragraphs]

        # Also extract from tables (skills/education tables are common in CVs)
        for table in doc.tables:
            for row in table.rows:
                paragraphs.extend(cell.text for cell in row.cells)

        text = "\n".join(paragraphs)
        return ExtractionResult(
            text=_normalise(text),
            parser_used="docx",
            page_count=0,       # python-docx doesn't expose page count
            is_image_based=False,
        )
    except Exception as exc:
        logger.error("DOCX extraction failed for %s: %s", path.name, exc)
        raise


# ──────────────────────────────────────────────────────────────────────────────
# Normalisation
# ──────────────────────────────────────────────────────────────────────────────


def _normalise(text: str) -> str:
    """Collapse whitespace and remove control characters."""
    text = re.sub(r"\r\n|\r", "\n", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[^\S\n]+$", "", text, flags=re.MULTILINE)
    return text.strip()
